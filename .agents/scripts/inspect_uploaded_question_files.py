from pathlib import Path
import fitz
from docx import Document

ROOT = Path("/home/runner/workspace")
ASSETS = ROOT / "attached_assets"
OUTPUT = ROOT / ".agents" / "outputs" / "uploaded-question-inspection"
OUTPUT.mkdir(parents=True, exist_ok=True)

pdfs = sorted(ASSETS.glob("*.pdf"))
for pdf_path in pdfs:
    doc = fitz.open(pdf_path)
    print(f"\n=== {pdf_path.name} ===")
    print(f"pages={doc.page_count} metadata={doc.metadata}")
    text_parts: list[str] = []
    for index, page in enumerate(doc):
        text = page.get_text("text")
        text_parts.append(f"\n--- PAGE {index + 1} ---\n{text}")
        if index in {0, min(1, doc.page_count - 1), doc.page_count - 1}:
            pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
            pix.save(OUTPUT / f"{pdf_path.stem}-page-{index + 1}.png")
    full_text = "".join(text_parts)
    (OUTPUT / f"{pdf_path.stem}.txt").write_text(full_text, encoding="utf-8")
    print(f"text_chars={len(full_text)} renders={min(doc.page_count, 2) + (1 if doc.page_count > 2 else 0)}")
    print(full_text[:3000])

docx_path = ASSETS / "New_Microsoft_Word_Document_(2)_1787047718026.docx"
doc = Document(docx_path)
print(f"\n=== {docx_path.name} ===")
print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)}")
docx_lines: list[str] = []
for paragraph in doc.paragraphs:
    if paragraph.text.strip():
        docx_lines.append(paragraph.text)
for table_index, table in enumerate(doc.tables, start=1):
    docx_lines.append(f"\n--- TABLE {table_index} ---")
    for row in table.rows:
        docx_lines.append(" | ".join(cell.text.replace("\n", " / ") for cell in row.cells))
docx_text = "\n".join(docx_lines)
(OUTPUT / f"{docx_path.stem}.txt").write_text(docx_text, encoding="utf-8")
print(docx_text[:5000])